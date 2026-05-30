# study_guide_linux_security.py

from docx import Document

doc = Document()
doc.add_heading('Linux OS Security Study Guide', 0)

# Chapter 1
doc.add_heading('Chapter 1: Security Threats to Linux', level=1)
doc.add_heading('Summary', level=2)
doc.add_paragraph(
    "This chapter covers the history and context of Linux security, "
    "the importance of ongoing security practices, the types of hackers, "
    "the open source security model, the C-I-A triad, and the origins and distributions of Linux."
)
doc.add_heading('Key Concepts & Terms', level=2)
doc.add_paragraph(
    "Apache, Application server, Availability, Bastion host, Black-hat hacker, "
    "Confidentiality, Debian, Distribution, Fedora, Firewall, GPL, GNU, Gray-hat hacker, "
    "Integrity, Kernel, Linux, Mint, Open source, RedHat, SSL, Snort, TLS, Ubuntu, Unix, White-hat hacker"
)
doc.add_heading('Assessment Questions', level=2)
doc.add_paragraph(
    "1. Which concept is part of the C-I-A triad?\n"
    "2. What makes up the core of the Linux OS?\n"
    "3. Which is an open source license?\n"
    "4. Security advantage of open source software?\n"
    "5. What % of Internet Web servers use Linux?\n"
    "6. Tool for intrusion-detection services?\n"
    "7. Open source license for GNU project?\n"
    "8. Linux is based on what?\n"
    "9. Primary purpose of hardening an OS?\n"
    "10. Risk of hand-building a kernel?\n"
    "11. What does C-I-A stand for?"
)

# Chapter 2
doc.add_heading('Chapter 2: Basic Components of Linux Security', level=1)
doc.add_heading('Summary', level=2)
doc.add_paragraph(
    "This chapter discusses the Linux kernel, boot process security, file permissions, "
    "virtualization, user authentication, service management, network security, "
    "security updates, and differences between Linux distributions."
)
doc.add_heading('Key Concepts & Terms', level=2)
doc.add_paragraph(
    "Access control lists (ACLs), BIND, Binary kernel, CentOS, CUPS, DNS, FTP, MySQL, "
    "NTP, Postfix, TFTP, and more."
)
doc.add_heading('Assessment Questions', level=2)
doc.add_paragraph(
    "1. Structure of the Linux kernel\n"
    "2. Linux Kernel Organization website\n"
    "3. Live CD distribution facts\n"
    "4. LILO boot loader security risks\n"
    "5. Services on a bastion host\n"
    "6. Linux GUI security issues\n"
    "7. Local authentication tools\n"
    "8. Discretionary access controls\n"
    "9. Tools that block IP addresses\n"
    "10. Role of mandatory access controls\n"
    "11. SSH client packages\n"
    "12. Kernel boot issues and recovery\n"
    "13. Controlling package updates\n"
    "14. Open source SMTP email options"
)

doc.save('Linux_OS_Security_Study_Guide.docx')
print("Word document 'Linux_OS_Security_Study_Guide.docx' created!")
